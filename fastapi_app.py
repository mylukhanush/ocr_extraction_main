from fastapi import FastAPI, UploadFile, File, Depends,Request, Header, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
import uvicorn
import socket
from io import BytesIO
from PIL import Image
import fitz
import asyncio
from contextlib import asynccontextmanager
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.middleware import SlowAPIMiddleware
from slowapi.errors import RateLimitExceeded
import tempfile
from docx import Document
import os
import subprocess
from dotenv import load_dotenv
from typing import Optional
from api_functions import setup_model_pool, pass_ocr_extraction,visa_ocr_extraction,eid_ocr_extraction,dl_ocr_extraction , e_visa_extraction ,get_medical_fitness_data, get_eid_application_details , mol_extraction ,get_status_change_data,get_insurance_card_details

# Implementation of Token Verification
load_dotenv()
SECRET_TOKEN = os.getenv("SECRET_TOKEN")

async def verify_token(Authorization: Optional[str] = Header(None)):

    print(f"Received Authorization Header: {Authorization}") # For debugging

    if Authorization is None:
        raise HTTPException(
            status_code=401,
            detail="Authorization header is missing",
        )
    token = Authorization # Since we are expecting a raw token

    if token != SECRET_TOKEN:
        raise HTTPException(
            status_code=403,
            detail="Wrong token. Access Denied",
        )


# Initialize FastAPI app with rate limiter
limiter = Limiter(key_func=get_remote_address)

@asynccontextmanager
async def lifespan(app: FastAPI):
    setup_model_pool()  # ✅ Load your model pool once
    yield
    # (Optional cleanup logic)


# Initialize FastAPI app
app1 = FastAPI(lifespan=lifespan)
# Attach the limiter to the app state
app1.state.limiter = limiter

app1.add_middleware(SlowAPIMiddleware)  # Add SlowAPI middleware

# Serve frontend static assets
app1.mount("/static", StaticFiles(directory="frontend/static"), name="static")

@app1.get("/", include_in_schema=False)
async def serve_frontend():
    return FileResponse("frontend/index.html")

@app1.get("/config", include_in_schema=False)
async def frontend_config():
    return JSONResponse({"token": SECRET_TOKEN})

# Handle Rate Limit Exceeded Exception
@app1.exception_handler(RateLimitExceeded)
async def rate_limit_handler(request, exc):
    return JSONResponse(content={"error": "Too many requests. Please try again later."}, status_code=429)


@app1.get("/whoami")
def whoami():
    return {"hostname": socket.gethostname()}


# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_passport_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_pass_details(request: Request,image: UploadFile = File(...), _: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))


        # elif file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await pass_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))


        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        # Get OCR extraction result from the imported function
        data, status_code = await pass_ocr_extraction(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            # Return the result as a JSON response
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400}
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_visa_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_pvisa_details(request: Request,image: UploadFile = File(...), _: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))


        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))


        # elif file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await visa_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))
        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        # Get OCR extraction result from the imported function
        data, status_code = await visa_ocr_extraction(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            # Return the result as a JSON response
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400}
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_eid_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_emiratesid_details(request: Request,image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))


        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))


        # elif file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await eid_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))
        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        # Get OCR extraction result from the imported function
        data, status_code = await eid_ocr_extraction(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            # Return the result as a JSON response
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400}
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_dl_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_driving_license_details(request: Request,image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))


        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))


        # elif file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await dl_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))
        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        # Get OCR extraction result from the imported function
        data, status_code = await dl_ocr_extraction(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            # Return the result as a JSON response
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400}
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


# Testing the FastAPI app locally
@app1.post("/extract_evisa_details")
@limiter.limit("5/minute") # Limit to 5 requests per minute
async def extract_e_visa_details(request: Request, image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if not image:
            return JSONResponse(content={"error": "No file part in the request"}, status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"}, status_code=400)

        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()

            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))

        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        data, status_code = await e_visa_extraction(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            # Return the result as a JSON response
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400
            }
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
            return JSONResponse(content={"error": str(e)}, status_code=500)


@app1.post("/extract_medical_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def medical_fitness_details_extraction(request: Request,image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file and get binary code/bytecode
            image_data = await image.read()
            #It takes raw byte code and turns into a image
            image = Image.open(BytesIO(image_data))

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))

        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        data, status_code = await get_medical_fitness_data(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400
            }
            return JSONResponse(content=result, status_code=status_code)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app1.post("/extract_eid_application_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_eid_application_details(request: Request, image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if not image:
            return JSONResponse(content={"error": "No file part in the request"}, status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"}, status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file and get binary code/bytecode
            image_data = await image.read()
            # It takes raw byte code and turns into a image
            image = Image.open(BytesIO(image_data))

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))
        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},
                                status_code=400)


        data, status_code = await get_eid_application_details(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            # Return the result as a JSON response
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400
            }
            return JSONResponse(content=result, status_code=status_code)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app1.post("/extract_mol_mb_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_mol_details(request: Request, image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if not image:
            return JSONResponse(content={"error": "No file part in the request"}, status_code=400)

        if image.filename == '' :
            return JSONResponse(content={"error": "No selected file"}, status_code=400)

        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file and get binary code/bytecode
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")   # Opens the PDF in memory
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)     # Create a transformation matrix to increase the resolution
            pix = page.get_pixmap(matrix=mat)     # Convert the page into a pixmap (image)
            img_data = pix.tobytes("png")         # Convert pixmap to  to PNG bytes
            image = Image.open(BytesIO(img_data)) # Opens the image from PNG bytes

        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg', 'jpeg', 'png' or 'pdf'."}, status_code=400)

        # Calling  an external function to extract MOL data from the image
        data, status_code = await mol_extraction(image)

        if status_code == 200:
            mol_number_start = data["mol_number"].lower().startswith("mb")
            if not mol_number_start:
                result = {
                    "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                    "sts": 400
                }
                return JSONResponse(content=result, status_code=status_code)

            result = {"data": data, "sts": 200, "msg": "Success"}
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400
            }
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app1.post("/extract_mol_st_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_mol_st_details(request: Request, image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if not image:
            return JSONResponse(content={"error": "No file part in the request"}, status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"}, status_code=400)

        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file and get binary code/bytecode
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")   # Opens the PDF in memory
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)     # Create a transformation matrix to increase the resolution
            pix = page.get_pixmap(matrix=mat)     # Convert the page into a pixmap (image)
            img_data = pix.tobytes("png")         # Convert pixmap to PNG bytes
            image = Image.open(BytesIO(img_data)) # Opens the image from PNG bytes

        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg', 'jpeg', 'png' or 'pdf'."}, status_code=400)

        # Calling  an external function to extract MOL data from the image
        data, status_code = await mol_extraction(image)

        if status_code == 200:
            mol_number_start = data["mol_number"].lower().startswith("st")
            if not mol_number_start:
                result = {
                    "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                    "sts": 400
                }
                return JSONResponse(content=result, status_code=status_code)

            result = {"data": data, "sts": 200, "msg": "Success"}
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400
            }
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app1.post("/extract_change_status_details")
@limiter.limit("5/minute")  # Limit to 5 requests per minute
async def extract_status_change_details(request: Request,image: UploadFile = File(...),_: None = Depends(verify_token)):
    try:
        await asyncio.sleep(1)
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))


        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))


        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        # Calling  an external function to extract MOL data from the image
        data, status_code = await get_status_change_data(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400
            }
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app1.post("/extract_insurance_details")
@limiter.limit("5/minute") # Limit to 5 requests per minute
async def extract_insurance_card_details(request: Request,image: UploadFile = File(...),_: None = Depends(verify_token)):
    print("Rate limit key:", get_remote_address(request))
    try:
        await asyncio.sleep(1)
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))


        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))


        else:
            return JSONResponse(content={"error": "Input file format must be 'jpg','jpeg','png' or 'pdf'."},status_code=400)

        # Calling  an external function to extract MOL data from the image
        data, status_code = await get_insurance_card_details(image)
        if status_code == 200:
            result = {"data": data, "sts": 200, "msg": "Success"}
            return JSONResponse(content=result, status_code=status_code)
        elif status_code == 400:
            result = {
                "msg": "The provided document is not suitable, so the requested details could not be extracted.",
                "sts": 400
            }
            return JSONResponse(content=result, status_code=status_code)

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


if __name__ == "__main__":
    # Run the FastAPI application using uvicorn
    # uvicorn.run(app, host="127.0.0.1", port=8000)
    uvicorn.run(app1, host="0.0.0.0", port=8000)



